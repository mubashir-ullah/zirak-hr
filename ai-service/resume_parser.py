import fitz  # PyMuPDF
import docx
import re
import openai
import os
from dotenv import load_dotenv
import json
from typing import Dict, List, Any, Optional

# Load environment variables
load_dotenv()

# Initialize OpenAI
openai.api_key = os.getenv("OPENAI_API_KEY")

class ResumeParser:
    """
    A class to parse resumes in PDF or DOCX format and extract structured information
    using OpenAI's GPT models.
    """
    
    def __init__(self, model="gpt-4"):
        """
        Initialize the ResumeParser with the specified OpenAI model.
        
        Args:
            model (str): The OpenAI model to use for parsing
        """
        self.model = model
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text from the PDF
        """
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def extract_text_from_pdf_bytes(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF bytes.
        
        Args:
            file_bytes (bytes): PDF file content as bytes
            
        Returns:
            str: Extracted text from the PDF
        """
        try:
            text = ""
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from PDF bytes: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text from the DOCX
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def extract_text_from_docx_bytes(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.
        
        Args:
            file_bytes (bytes): DOCX file content as bytes
            
        Returns:
            str: Extracted text from the DOCX
        """
        try:
            from io import BytesIO
            doc = docx.Document(BytesIO(file_bytes))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX bytes: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text from the file
        """
        if file_path.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")
    
    def extract_text_from_bytes(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from file bytes based on the file type.
        
        Args:
            file_bytes (bytes): File content as bytes
            file_type (str): File type ('pdf' or 'docx')
            
        Returns:
            str: Extracted text from the file
        """
        if file_type.lower() == 'pdf':
            return self.extract_text_from_pdf_bytes(file_bytes)
        elif file_type.lower() == 'docx':
            return self.extract_text_from_docx_bytes(file_bytes)
        else:
            raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")
    
    def parse_with_gpt(self, text: str) -> Dict[str, Any]:
        """
        Parse resume text using OpenAI's GPT model.
        
        Args:
            text (str): Resume text to parse
            
        Returns:
            Dict[str, Any]: Structured resume data
        """
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": """You are a resume parsing assistant. Extract the following information from the resume text:
                    - Full Name
                    - Email
                    - Phone
                    - Skills (as a list)
                    - Years of Experience
                    - Education (as a list of objects with degree, institution, and year)
                    - Work Experience (as a list of objects with title, company, duration, and description)
                    - Country
                    - City
                    - Languages (as a list with proficiency levels)
                    - LinkedIn URL
                    - GitHub URL
                    - Portfolio URL
                    
                    For German language specifically, indicate the level (None, A1, A2, B1, B2, C1, C2, or Native).
                    
                    Format your response as a valid JSON object with these keys: fullName, email, phone, skills, experience, education, workExperience, country, city, languages, germanLevel, linkedinUrl, githubUrl, portfolioUrl.
                    For skills, return an array of strings. For experience, return the number of years as a string."""},
                    {"role": "user", "content": text}
                ],
                temperature=0.3,
                max_tokens=1500
            )
            
            parsed_data = json.loads(response.choices[0].message.content)
            return parsed_data
        except Exception as e:
            raise Exception(f"Error parsing resume with GPT: {str(e)}")
    
    def suggest_skills(self, resume_text: str, current_skills: List[str]) -> Dict[str, Any]:
        """
        Suggest additional skills based on resume text and current skills.
        
        Args:
            resume_text (str): Resume text
            current_skills (List[str]): List of current skills
            
        Returns:
            Dict[str, Any]: Dictionary with suggested skills and confidence scores
        """
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": """You are a career advisor specialized in tech skills. 
                    Based on the resume text and current skills, suggest additional relevant skills that would enhance the candidate's profile.
                    Focus on technical skills, tools, frameworks, and methodologies that are in demand in the job market.
                    Return only a JSON array of strings with the suggested skills. Do not include any of the current skills in your suggestions."""},
                    {"role": "user", "content": f"Resume text: {resume_text}\n\nCurrent skills: {', '.join(current_skills)}"}
                ],
                temperature=0.5,
                max_tokens=500
            )
            
            suggested_skills = json.loads(response.choices[0].message.content)
            
            # Calculate confidence scores (simplified)
            confidence_scores = [0.9 - (i * 0.05) for i in range(len(suggested_skills))]
            confidence_scores = [max(score, 0.5) for score in confidence_scores]  # Ensure minimum 0.5 confidence
            
            return {
                "skills": suggested_skills,
                "confidence": confidence_scores
            }
        except Exception as e:
            raise Exception(f"Error suggesting skills: {str(e)}")
    
    def generate_resume_summary(self, parsed_data: Dict[str, Any]) -> str:
        """
        Generate a professional summary based on parsed resume data.
        
        Args:
            parsed_data (Dict[str, Any]): Parsed resume data
            
        Returns:
            str: Professional summary
        """
        try:
            # Create a prompt with the parsed data
            prompt = f"""
            Name: {parsed_data.get('fullName', 'N/A')}
            Experience: {parsed_data.get('experience', 'N/A')} years
            Skills: {', '.join(parsed_data.get('skills', []))}
            Education: {', '.join([f"{edu.get('degree')} from {edu.get('institution')}" for edu in parsed_data.get('education', [])])}
            Work Experience: {', '.join([f"{exp.get('title')} at {exp.get('company')}" for exp in parsed_data.get('workExperience', [])])}
            """
            
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": """You are a professional resume writer. 
                    Create a concise, impactful professional summary (2-3 sentences) based on the candidate's information.
                    The summary should highlight their experience, key skills, and unique value proposition.
                    Write in first person and focus on achievements and expertise."""},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=200
            )
            
            return response.choices[0].message.content.strip()
        except Exception as e:
            raise Exception(f"Error generating resume summary: {str(e)}")
    
    def analyze_resume_for_job(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """
        Analyze how well a resume matches a job description.
        
        Args:
            resume_text (str): Resume text
            job_description (str): Job description text
            
        Returns:
            Dict[str, Any]: Analysis results including match score and recommendations
        """
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": """You are a career coach and resume expert.
                    Analyze how well the candidate's resume matches the job description.
                    Provide:
                    1. An overall match score (0-100)
                    2. Matching skills (skills in the resume that match the job requirements)
                    3. Missing skills (skills in the job description that are not in the resume)
                    4. Recommendations to improve the resume for this specific job
                    
                    Format your response as a JSON object with these keys: matchScore, matchingSkills, missingSkills, recommendations."""},
                    {"role": "user", "content": f"Resume:\n{resume_text}\n\nJob Description:\n{job_description}"}
                ],
                temperature=0.5,
                max_tokens=1000
            )
            
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            raise Exception(f"Error analyzing resume for job: {str(e)}")

# Example usage
if __name__ == "__main__":
    parser = ResumeParser()
    
    # Example: Parse a resume file
    # resume_text = parser.extract_text("path/to/resume.pdf")
    # parsed_data = parser.parse_with_gpt(resume_text)
    # print(json.dumps(parsed_data, indent=2))
